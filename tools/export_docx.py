#!/usr/bin/env python3
"""
Génère public/downloads/ins-mv-rogue-srd.docx à partir des fichiers markdown du SRD.

Dépendances :
    pip install python-docx

Usage :
    python tools/export_docx.py
"""

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path("src/content/docs")
OUTPUT   = Path("public/downloads/ins-mv-rogue-srd.docx")

EXPORT_ORDER = [
    "index",
    "contexte/contexte",
    "personnage/caracteristiques",
    "personnage/creation",
    "personnage/progression",
    "personnage/rang",
    "personnage/reincarnation",
    "mecanique/resolution",
    "mecanique/combat",
    "mecanique/competences",
    "mecanique/energie",
    "mecanique/pouvoirs",
    "mecanique/blessures",
    "reference/mots-cles",
    "reference/etats",
    "reference/equipement/index",
    "reference/equipement/armes-feu",
    "reference/equipement/melee",
    "reference/equipement/distance",
    "reference/equipement/explosifs",
    "reference/equipement/protections",
    "reference/equipement/boucliers",
]

# ── Couleurs thème ────────────────────────────────────────────────────────────

BLUE_ANGEL  = RGBColor(0x2B, 0x6C, 0xB0)   # saphir angélique
RED_INTENSE = RGBColor(0xC5, 0x30, 0x30)   # rouge intensité
GOLD_DIVINE = RGBColor(0x92, 0x70, 0x0A)   # or divin
GREY_DIM    = RGBColor(0x4A, 0x4A, 0x6A)

INDENT_STEP = Cm(0.5)
TABLE_PLACEHOLDER_RE = re.compile(r'\x00TABLE:(\d+)\x00')

# ── Helpers XML ──────────────────────────────────────────────────────────────

def _set_cell_shading(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    tcPr.append(shd)


def _add_border_left(paragraph, hex_color="2B6CB0", size_pt=12):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(size_pt * 8))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _apply_indent(paragraph, indent, bordered=True):
    if indent is None:
        return
    paragraph.paragraph_format.left_indent = indent
    if bordered:
        _add_border_left(paragraph, size_pt=8)


# ── Markdown inline parser ────────────────────────────────────────────────────

_INLINE_RE = re.compile(
    r'(\*\*\*(?P<boi>.+?)\*\*\*'
    r'|\*\*(?P<bo>.+?)\*\*'
    r'|\*(?P<it>.+?)\*'
    r'|`(?P<code>.+?)`'
    r'|\[(?P<lt>[^\]]+)\]\([^\)]+\)'   # lien → texte seul (mise en forme imbriquée gérée séparément)
    r')'
)


def _strip_link_inline_markup(text):
    """Retire un éventuel gras/italique enveloppant tout le texte d'un lien, renvoie (texte, gras, italique)."""
    m = re.match(r'^\*\*\*(.+)\*\*\*$', text, re.DOTALL)
    if m:
        return m.group(1), True, True
    m = re.match(r'^\*\*(.+)\*\*$', text, re.DOTALL)
    if m:
        return m.group(1), True, False
    m = re.match(r'^\*(.+)\*$', text, re.DOTALL)
    if m:
        return m.group(1), False, True
    return text, False, False


def _html_inline_to_md(text):
    """Convertit les balises HTML inline résiduelles (strong/em/br) en équivalents markdown."""
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<strong>(.*?)</strong>', r'**\1**', text, flags=re.DOTALL)
    text = re.sub(r'<em>(.*?)</em>', r'*\1*', text, flags=re.DOTALL)
    return text


def add_inline(para, raw_text, bold=False, italic=False):
    """Ajoute du texte avec mise en forme inline (gras/italique/code/liens), imbrications comprises."""
    text = _html_inline_to_md(raw_text)
    text = re.sub(r'<[^>]+>', '', text)   # strip HTML résiduel
    last = 0
    for m in _INLINE_RE.finditer(text):
        before = text[last:m.start()]
        if before:
            r = para.add_run(before)
            r.bold = bold; r.italic = italic
        if m.group("boi"):
            add_inline(para, m.group("boi"), bold=True, italic=True)
        elif m.group("bo"):
            add_inline(para, m.group("bo"), bold=True, italic=italic)
        elif m.group("it"):
            add_inline(para, m.group("it"), bold=bold, italic=True)
        elif m.group("code"):
            r = para.add_run(m.group("code"))
            r.font.name = "Courier New"; r.font.size = Pt(9)
            r.bold = bold; r.italic = italic
        elif m.group("lt"):
            inner, lbold, litalic = _strip_link_inline_markup(m.group("lt"))
            r = para.add_run(inner)
            r.font.color.rgb = BLUE_ANGEL
            r.bold = bold or lbold
            r.italic = italic or litalic
        last = m.end()
    tail = text[last:]
    if tail:
        r = para.add_run(tail)
        r.bold = bold; r.italic = italic


# ── Frontmatter ───────────────────────────────────────────────────────────────

def strip_frontmatter(text):
    """Retourne (title, body) après suppression du frontmatter YAML."""
    m = re.match(r'^---\n(.*?)\n---\n?(.*)', text, re.DOTALL)
    if not m:
        return '', text
    fm, body = m.group(1), m.group(2)
    t = re.search(r'^title:\s*["\']?(.+?)["\']?\s*$', fm, re.MULTILINE)
    return (t.group(1).strip() if t else ''), body.strip()


# ── Extraction des tableaux HTML (statiques ou générés en JS) ─────────────────

def _clean_cell_html(html):
    text = _html_inline_to_md(html)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # lien -> texte
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _parse_html_table(table_html, scripts):
    headers = [_clean_cell_html(re.sub(r'\s*⇅\s*$', '', h))
               for h in re.findall(r'<th\b[^>]*>(.*?)</th>', table_html, re.DOTALL)]

    rows = []
    for tr in re.findall(r'<tr\b[^>]*>(.*?)</tr>', table_html, re.DOTALL):
        if '<th' in tr:
            continue
        cells = re.findall(r'<td\b[^>]*>(.*?)</td>', tr, re.DOTALL)
        if not cells:
            continue
        rows.append([_clean_cell_html(c) for c in cells])

    # Tableau vide côté HTML : généré dynamiquement par un <script> (recherche/filtres JS)
    if not rows and scripts:
        for script in scripts:
            data_m = re.search(r'var\s+DATA_\w+\s*=\s*(\[.*?\]);', script, re.DOTALL)
            map_m  = re.search(r'\.map\(\s*function\s*\(r\)\s*\{(.*?)\}\)\s*\.join', script, re.DOTALL)
            if not data_m or not map_m:
                continue
            try:
                data = json.loads(data_m.group(1))
            except json.JSONDecodeError:
                continue
            # Une clé par cellule <td> du gabarit (dernière référence r[...] rencontrée dans
            # son segment), pour ignorer les éventuelles variables intermédiaires (ex: couleur
            # calculée en amont) sans dépendre de la façon dont la cellule est concaténée.
            keys = []
            for seg in re.split(r'(?=<td)', map_m.group(1)):
                if '<td' not in seg:
                    continue
                seg_keys = re.findall(r"r\[([\"'])((?:(?!\1).)*)\1\]", seg)
                if seg_keys:
                    keys.append(seg_keys[-1][1])
            if not data or len(keys) != len(headers):
                continue
            for row in data:
                rows.append([str(row.get(k, '') if row.get(k) is not None else '') for k in keys])
            break

    return {"headers": headers, "rows": rows}


def extract_tables(text):
    """Repère les tableaux (statiques ou générés en JS via <script>), les remplace par un
    marqueur, et renvoie (texte_modifié, liste_de_tableaux)."""
    tables = []

    # Bloc complet "widget de recherche + table vide + script de rendu"
    def repl_widget(m):
        table_html = m.group(1)
        script     = m.group(2)
        tables.append(_parse_html_table(table_html, [script]))
        return f'\n\x00TABLE:{len(tables) - 1}\x00\n'

    text = re.sub(
        r'<div id="table-[\w-]+">(.*?)</div>\s*\n<script>(.*?)</script>',
        repl_widget, text, flags=re.DOTALL,
    )

    # Tableaux HTML statiques restants
    def repl_static(m):
        tables.append(_parse_html_table(m.group(0), []))
        return f'\n\x00TABLE:{len(tables) - 1}\x00\n'

    text = re.sub(r'<table\b.*?</table>', repl_static, text, flags=re.DOTALL)

    # Scripts orphelins restants (aucun tableau associé trouvé) : à ignorer totalement
    text = re.sub(r'<script>.*?</script>', '', text, flags=re.DOTALL)

    return text, tables


def render_table(doc, table, indent=None):
    headers, rows = table["headers"], table["rows"]
    if not rows:
        return
    cols = len(headers) if headers else max(len(r) for r in rows)
    t = doc.add_table(rows=(1 if headers else 0) + len(rows), cols=cols)
    t.style = "Table Grid"

    ri = 0
    if headers:
        for ci, h in enumerate(headers):
            cell = t.rows[0].cells[ci]
            cell.text = h
            _set_cell_shading(cell, "E8E8F0")
            for run in cell.paragraphs[0].runs:
                run.bold = True
        ri = 1

    for row in rows:
        for ci in range(cols):
            val = row[ci] if ci < len(row) else ''
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
        ri += 1

    doc.add_paragraph()


# ── Convertisseur principal ───────────────────────────────────────────────────

def process_content(doc, content, tables, indent=None):
    lines = content.split('\n')
    i = 0
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        render_table(doc, {"headers": table_rows[0], "rows": table_rows[1:]}, indent=indent)
        table_rows = []

    def add_para(style=None):
        p = doc.add_paragraph(style=style)
        _apply_indent(p, indent, bordered=False)
        return p

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Ligne vide
        if not stripped:
            flush_table()
            i += 1
            continue

        # Marqueur de tableau extrait en amont
        tm = TABLE_PLACEHOLDER_RE.match(stripped)
        if tm:
            flush_table()
            render_table(doc, tables[int(tm.group(1))], indent=indent)
            i += 1
            continue

        # Bloc <div ...> ... </div> ou <a ...> ... </a> (admonitions, cartes-liens, etc.)
        m_block = re.match(r'^<(div|a)\b[^>]*>', stripped)
        if m_block and not stripped.startswith('</'):
            flush_table()
            tag = m_block.group(1)
            class_m = re.search(r'class="([^"]*)"', stripped)
            classes = class_m.group(1).split() if class_m else []

            open_re, close_re = f'<{tag}', f'</{tag}'
            depth = stripped.count(open_re) - stripped.count(close_re)
            inner_lines = []
            i += 1
            while i < len(lines) and depth > 0:
                l = lines[i]
                depth += l.count(open_re) - l.count(close_re)
                if depth <= 0:
                    l = re.sub(rf'</{tag}>\s*$', '', l)
                    if l.strip():
                        inner_lines.append(l)
                    i += 1
                    break
                inner_lines.append(l)
                i += 1

            render_div_block(doc, classes, '\n'.join(inner_lines), tables, indent)
            continue

        if stripped.startswith('</div') or stripped.startswith('</a'):
            i += 1
            continue

        # Séparateur horizontal
        if re.match(r'^[-*_]{3,}$', stripped):
            flush_table()
            p = add_para()
            r = p.add_run('─' * 60)
            r.font.color.rgb = GREY_DIM
            i += 1
            continue

        # Titres
        m = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if m:
            flush_table()
            level = min(len(m.group(1)), 6)
            h = doc.add_heading('', level=level)
            _apply_indent(h, indent, bordered=False)
            add_inline(h, m.group(2))
            i += 1
            continue

        # Tableau markdown
        if '|' in stripped and stripped.startswith('|'):
            if re.match(r'^\|[-:\s|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            table_rows.append([_clean_cell_html(c) for c in cells])
            i += 1
            continue

        flush_table()

        # Listes à puces
        m_ul = re.match(r'^(\s*)-\s+(.+)$', line)
        if m_ul:
            lvl = len(m_ul.group(1)) // 2
            p = add_para(style='List Bullet')
            p.paragraph_format.left_indent = (indent or Cm(0)) + Cm(0.5 + lvl * 0.5)
            add_inline(p, m_ul.group(2))
            i += 1
            continue

        # Listes numérotées
        m_ol = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if m_ol:
            lvl = len(m_ol.group(1)) // 2
            p = add_para(style='List Number')
            p.paragraph_format.left_indent = (indent or Cm(0)) + Cm(0.5 + lvl * 0.5)
            add_inline(p, m_ol.group(2))
            i += 1
            continue

        # Paragraphe ordinaire
        p = add_para()
        add_inline(p, stripped)
        i += 1

    flush_table()


def render_div_block(doc, classes, inner_text, tables, indent):
    """Rend un bloc <div> (admonition, carte, grille de cartes...) en isolant un éventuel
    titre puis en traitant le corps via process_content (récursif, avec indentation)."""

    # Grille de cartes : chaque enfant est son propre bloc, pas de titre au niveau de la grille
    if 'rogue-card-grid' in classes:
        process_content(doc, inner_text, tables, indent=indent)
        return

    title = None

    m = re.search(r'<p class="admonition-title">(.*?)</p>', inner_text, re.DOTALL)
    if m:
        title = _clean_cell_html(m.group(1))
        inner_text = inner_text[:m.start()] + inner_text[m.end():]
    else:
        m2 = re.match(r'\s*<strong>(.*?)</strong>\s*(<br\s*/?>)?', inner_text, re.DOTALL)
        if m2:
            title = _clean_cell_html(m2.group(1))
            inner_text = inner_text[m2.end():]

    body_indent = (indent or Cm(0)) + INDENT_STEP

    if title:
        p = doc.add_paragraph()
        _apply_indent(p, indent, bordered=True)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.color.rgb = BLUE_ANGEL
    else:
        body_indent = indent

    body = inner_text.strip('\n')
    if body.strip():
        process_content(doc, body, tables, indent=body_indent)


# ── Document setup ────────────────────────────────────────────────────────────

def setup_styles(doc):
    """Applique les styles de base au document."""
    normal = doc.styles['Normal']
    normal.font.name = 'Cambria'
    normal.font.size = Pt(11)

    for i in range(1, 7):
        try:
            h = doc.styles[f'Heading {i}']
            h.font.name  = 'Cambria'
            h.font.color.rgb = BLUE_ANGEL if i <= 2 else RGBColor(0x1A, 0x1A, 0x2E)
            h.font.size  = Pt(max(20 - i * 2, 10))
        except KeyError:
            pass

    # Marges du document
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(3)


# ── Page de titre ─────────────────────────────────────────────────────────────

def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("ROGUE")
    r.font.name  = "Cambria"
    r.font.size  = Pt(32)
    r.font.bold  = True
    r.font.color.rgb = BLUE_ANGEL

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("System Reference Document")
    r2.font.name = "Cambria"
    r2.font.size = Pt(18)
    r2.font.color.rgb = GREY_DIM

    doc.add_page_break()


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_styles(doc)
    add_title_page(doc)

    for slug in EXPORT_ORDER:
        path = DOCS_DIR / (slug + '.md')
        if not path.exists():
            print(f"  ⚠ {path} introuvable, ignoré.")
            continue

        text         = path.read_text(encoding='utf-8')
        title, body  = strip_frontmatter(text)
        body, tables = extract_tables(body)
        print(f"  ✓ {slug}: {title}")

        if title:
            doc.add_heading(title, level=1)

        process_content(doc, body, tables)
        doc.add_page_break()

    doc.save(OUTPUT)
    print(f"\nFichier généré : {OUTPUT}")


if __name__ == '__main__':
    main()
